# -*- coding: utf-8 -*-

from tkinter import *
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from tkinter import filedialog as fd
import os


if __name__ == "__main__":
    techpath = 'tech.txt'
    grouppath = "group.txt"
    peoplepath = 'people.txt'

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)


def textfieldtodocx(event):
    line = textField.get(2.0, END)
    mylist = line.split('\n')
    mylist.pop()
    for i in range(len(mylist)):
        mylist[i] = str(i+1)+'|'+str(mylist[i])
    records = []
    for x in range(len(mylist)):
        tmplist = mylist[x].split('|')
        tmplist[2] = tmplist[2] + '\n(' + tmplist[3] + ')'
        del tmplist[3]
        records.append(tmplist)
    actpar = document.add_paragraph('Акт о переносе техники №' + actNumber.get())
    actpar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = document.add_table(rows=1, cols=5, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(0.4)      # № п/п
    table.columns[1].width = Inches(2.1)    # Перемещение
    table.columns[2].width = Inches(1.3)    # Инв. номер
    table.columns[3].width = Inches(1.3)    # Забрали из отдела\n ФИО
    table.columns[4].width = Inches(1.3)    # Поставили в отдел\n ФИО
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Перемещение'
    hdr_cells[2].text = 'Инв. номер'
    hdr_cells[3].text = 'Забрали из отдела\n ФИО'
    hdr_cells[4].text = 'Поставили в отдел\n ФИО'

    for num, ids, invent, out, inp in records:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = ids
        row_cells[2].text = invent
        row_cells[3].text = out
        row_cells[4].text = inp

    datepar = document.add_paragraph('\n\n' + dateEntry.get()+'			Подпись_____________		'
                                     + people.get())
    datepar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    file_name = fd.asksaveasfilename(initialfile=('#'+actNumber.get()+'_'+dateEntry.get()),
                                     defaultextension=".docx", filetypes=[('docx files', '*.docx')])
    document.save(file_name)


def nextvar(event):
    textField.config(state=NORMAL)
    textField.insert(END, '\n' + tech.get() + '|' + inventNumber.get() + '|' + codeNumber.get() + '|' +
                     groupOut.get() + '|' + groupIn.get())
    textField.config(state=DISABLED, bg="gainsboro")


def edittextfield(event):
    textField.config(state=NORMAL, bg="white")


def restart_program(event):
    """Restarts the current program.
    Note: this function does not return. Any cleanup action (like
    saving data) must be done before calling this function.
    https://stackoverflow.com/questions/41655618/restart-program-tkinter"""
    python = sys.executable
    os.execl(python, python, * sys.argv)


techtxt = open(techpath, 'r', encoding="utf-8")
techlist = techtxt.read().splitlines()
grouptxt = open(grouppath, 'r', encoding="utf-8")
grouplist = grouptxt.read().splitlines()
peopletxt = open(peoplepath, 'r', encoding="utf-8")
peoplelist = peopletxt.read().splitlines()

root = Tk()
root.title('Mover 9000')
root.resizable(False, False)

numberFrame = Frame(root)
numberFrame.pack(fill=X)
clearButton = Button(numberFrame, text='Очистить\nвсё', pady=10, command=exit, bd=3, relief=GROOVE)
clearButton.grid(row=0, column=0)
clearButton.bind('<ButtonRelease-1>', restart_program)
emptyLabel = Label(numberFrame, text=' ', padx=160)
emptyLabel.grid(row=0, column=1)
actNumber = StringVar(root)
actLabel = Label(numberFrame, text='Акт о переносе техники №', font='Arial 14')
actLabel.grid(row=0, column=2, sticky=E)
actNumberEntry = Entry(numberFrame, font='Arial 14', justify='center', textvariable=actNumber, width=4)
actNumberEntry.grid(row=0, column=3, sticky=W)

# ---------------------------BUTTONS-----------------------------------------------------------------------------------
buttonFrame = Frame(root)
buttonFrame.pack()
tech = StringVar(root)
tech.set(techlist[0])
groupOut = StringVar(root)
groupIn = StringVar(root)
groupOut.set(grouplist[0])
groupIn.set(grouplist[1])
inventNumber = StringVar(root)
codeNumber = StringVar(root)
techMenu = OptionMenu(buttonFrame, tech, *techlist)
techMenu.config(width=21, font='Arial 12', relief=GROOVE)
inventNumberEntry = Entry(buttonFrame, font='Arial 17', justify='center', textvariable=inventNumber, width=15)
codeNumberEntry = Entry(buttonFrame, font='Arial 17', justify='center', textvariable=codeNumber, width=10)
groupOutMenu = OptionMenu(buttonFrame, groupOut, *grouplist)
groupOutMenu.config(width=21, font='Arial 12', relief=GROOVE)
groupInMenu = OptionMenu(buttonFrame, groupIn, *grouplist)
groupInMenu.config(width=21, font='Arial 12', relief=GROOVE)

techMenu.pack(side=LEFT)
inventNumberEntry.pack(side=LEFT)
codeNumberEntry.pack(side=LEFT)
groupOutMenu.pack(side=LEFT)
groupInMenu.pack(side=LEFT)
# --------------------------BUTTONS--END-------------------------------------------------------------------------------

textFrame = Frame(root)
textFrame.pack()
buttonPlus = Button(textFrame, text='добавить', font='Arial 13', justify='center', relief=GROOVE)
buttonPlus.bind('<ButtonRelease-1>', nextvar)
buttonPlus.pack(side=TOP, fill=X)
textField = Text(textFrame, width=112, state=DISABLED, bg="gainsboro", font='Arial 12')
scrollText = Scrollbar(textFrame, command=textField.yview)
scrollText.pack(side=RIGHT, fill=Y)
textField.config(yscrollcommand=scrollText.set)
textField.pack()
doneFrame = Frame(root)
doneFrame.pack(fill=X)
editButton = Button(doneFrame, text='Редактировать список', font='Arial 10', bd=2, relief=GROOVE)
doneButton = Button(doneFrame, text='СОХРАНИТЬ', font='Arial 16', bd=3, relief=GROOVE)
editButton.pack(side=TOP, fill=X)
editButton.bind('<ButtonRelease-1>', edittextfield)
nLabel = Label(doneFrame, text='\n', font='Arial 2')
nLabel.pack(fill=X)
dateFrame = Frame(doneFrame)
dateFrame.pack(fill=X)

noow = datetime.now()
docDate = StringVar()
dateEntry = Entry(dateFrame, font='Arial 14', justify='center', textvariable=docDate)
dateEntry.insert(0, str(noow.day)+'.'+str(noow.month)+'.'+str(noow.year))
dateEntry.pack(side=LEFT)

people = StringVar(root)
people.set(peoplelist[0])
peopleMenu = OptionMenu(dateFrame, people, *peoplelist)
peopleMenu.config(width=21, font='Arial 12', relief=GROOVE)
peopleMenu.pack(side=RIGHT)

n2Label = Label(doneFrame, text='\n', font='Arial 2')
n2Label.pack(fill=X)
doneButton.pack(fill=X)
doneButton.bind('<ButtonRelease-1>', textfieldtodocx)
doneButton.bind('Control+s>', textfieldtodocx)
root.mainloop()
